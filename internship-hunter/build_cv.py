"""
Generates Harry_Winter_CV_2026.docx
Run: python build_cv.py
Every claim on this CV must survive a follow-up question at interview.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


ACCENT = RGBColor(0x1A, 0x1A, 0x6E)   # dark navy
MID    = RGBColor(0x44, 0x44, 0x44)   # dark grey for dates


def _rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = ACCENT
    _rule(doc)


def _entry(doc, title, org, dates):
    """Role/project header line: bold title | org  (right-aligned date)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(1)

    r_title = p.add_run(title)
    r_title.bold = True
    r_title.font.size = Pt(10.5)

    if org:
        r_org = p.add_run(f"  ·  {org}")
        r_org.font.size = Pt(10)
        r_org.font.color.rgb = MID

    # Right-aligned date via tab stop at right margin
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_ALIGN_PARAGRAPH.RIGHT)
    r_date = p.add_run(f"\t{dates}")
    r_date.font.size = Pt(9.5)
    r_date.font.color.rgb = MID


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.font.size = Pt(10)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin   = Inches(0.8)
        section.right_margin  = Inches(0.8)

    # ── Name ──────────────────────────────────────────────────────────────────
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(2)
    r = p_name.add_run("Harry Winter")
    r.bold = True
    r.font.size = Pt(20)

    # ── Contact ────────────────────────────────────────────────────────────────
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(8)
    cr = p_contact.add_run(
        "+44 7453419094  ·  harrywinter.uk@gmail.com  ·  London, UK\n"
        "linkedin.com/in/harry-winter-1472b5330"
    )
    cr.font.size = Pt(9.5)
    cr.font.color.rgb = MID

    # ── Education ──────────────────────────────────────────────────────────────
    _section(doc, "Education")
    _entry(doc, "BSc Information Management for Business", "University College London", "Expected June 2027")
    _bullet(doc, "Predicted First Class Honours")
    _bullet(doc, "Modules: Business Analytics  ·  Quantitative Methods for Business  ·  Programming  ·  Accounting for Decision Makers")
    _bullet(doc, "Runner-up, UCL × Svitlo Design Sprint — built a complete business solution for a real client: KPI definition, customer research, data analysis, financial justification; presented to senior stakeholders and judges")

    _entry(doc, "A-Levels: Mathematics (A), Business (A), Physics (B)", "Hitchin Boys' School", "2022 – 2024")

    # ── Projects ───────────────────────────────────────────────────────────────
    _section(doc, "Projects")

    _entry(doc, "Algorithmic Trading System", "Python  ·  data engineering  ·  automation", "2025 – Present")
    _bullet(doc, "Built a multi-component Python pipeline that ingests live financial data daily: macroeconomic indicators (FRED API), equity price history (yfinance), market sentiment from Reddit (PRAW), and news RSS feeds")
    _bullet(doc, "Runs autonomously on a weekday scheduler — data ingestion, feature processing, strategy evaluation, and execution flow without manual intervention")
    _bullet(doc, "Paper trading live on Alpaca Markets; strategies evaluated using out-of-sample validation to guard against overfitting to historical data")
    _bullet(doc, "Multi-phase architecture with distinct modules for data ingestion, feature engineering, strategy search, risk management, and execution — built and iterated across 16 development phases")

    _entry(doc, "Internship Hunter", "Python  ·  Anthropic API  ·  SQLite  ·  Streamlit", "2026")
    _bullet(doc, "Automated job discovery system: scrapes Greenhouse and Lever ATS boards across a curated company list, deduplicates by URL, filters by role relevance")
    _bullet(doc, "Cover letter and cold outreach generation via Claude Sonnet API — system prompt cached to minimise cost while maintaining quality")
    _bullet(doc, "Streamlit dashboard for full application lifecycle: status tracking, follow-up dates, cost monitoring")

    # ── Experience ─────────────────────────────────────────────────────────────
    _section(doc, "Experience")

    _entry(doc, "Catering Assistant", "Edible Food Design", "September 2025 – Present")
    _bullet(doc, "High-end event catering including private and royal events — consistent reliability and precision under pressure")

    # ── Academic Projects ──────────────────────────────────────────────────────
    _section(doc, "Academic Projects")

    _entry(doc, "IMB Design Sprint — Information Systems Design", "NatWest & JP Morgan  ·  UCL", "Oct – Nov 2025")
    _bullet(doc, "Led team designing a role-based information system prototype: user flows, UML class diagrams, use case models, application architecture")
    _bullet(doc, "Established requirements traceability from business objectives to technical implementation; delivered structured presentations on trade-offs and MVP strategy")

    _entry(doc, "Platform Analysis", "Pottermore  ·  UCL", "Sep – Dec 2024")
    _bullet(doc, "Mapped user journeys, stakeholder interactions, and value flows; developed operational KPIs to evaluate platform engagement and scalability")

    # ── Leadership ─────────────────────────────────────────────────────────────
    _section(doc, "Leadership & Activities")

    _entry(doc, "Transition Mentor", "UCL School of Management", "September 2025 – Present")
    _bullet(doc, "Selected to mentor 15 first-year students; deliver workshops, one-on-one guidance, and faculty feedback on programme outcomes")

    _entry(doc, "Member", "UCL Investment Society", "November 2024 – Present")
    _bullet(doc, "Equity research workshops; team stock pitch competition presenting full valuation to judges")

    # ── Skills & Additional ────────────────────────────────────────────────────
    _section(doc, "Skills & Additional")

    p_skills = doc.add_paragraph()
    p_skills.paragraph_format.space_before = Pt(5)
    p_skills.paragraph_format.space_after = Pt(3)
    r_lbl = p_skills.add_run("Technical: ")
    r_lbl.bold = True
    r_lbl.font.size = Pt(10)
    p_skills.add_run(
        "Python (pandas, numpy, yfinance, requests, sqlite3, Streamlit, Anthropic SDK)  ·  Excel financial modelling  ·  SQL  ·  Git"
    ).font.size = Pt(10)

    p_add = doc.add_paragraph()
    p_add.paragraph_format.space_before = Pt(2)
    r_lbl2 = p_add.add_run("Additional: ")
    r_lbl2.bold = True
    r_lbl2.font.size = Pt(10)
    p_add.add_run(
        "Snowboarder, runner, jujitsu, ice hockey  ·  Volunteer, PHASE Mental Health Charity  ·  "
        "Reads: Bloomberg Daybreak daily, Margin of Safety, Moneyball"
    ).font.size = Pt(10)

    # ── Save ───────────────────────────────────────────────────────────────────
    out = Path(r"C:\Users\Winte\OneDrive\CV STUFF\Harry_Winter_CV_2026.docx")
    doc.save(str(out))
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
