"""
Cover letter and outreach generator using Claude Sonnet with prompt caching.
System prompt (~profile) is cached — subsequent calls in the same session cost ~10% of the first.

Two outreach modes:
  - email: cold email with subject line, 50-75 words, addressed to named person
  - linkedin: direct message, 40-60 words, no subject line
"""

import anthropic
from anthropic.types import TextBlock
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from config import ANTHROPIC_API_KEY, GENERATE_MODEL, PROFILE_PATH, COVER_LETTER_DIR, OUTREACH_DIR

COVER_LETTER_DIR.mkdir(parents=True, exist_ok=True)
OUTREACH_DIR.mkdir(parents=True, exist_ok=True)

_profile_text = PROFILE_PATH.read_text(encoding="utf-8")

# ── System prompts ─────────────────────────────────────────────────────────────

COVER_LETTER_SYSTEM = f"""You write professional cover letters for Harry Winter, a UCL undergraduate applying for summer 2026 internships (June–September).

QUALITY STANDARD: A senior recruiter reads 200 CVs a day. This letter must earn its place in the 10% that get a response. Mediocre is failure.

STRUCTURE (4 paragraphs, under 350 words total):
1. Opening: lead with the single strongest match between Harry's background and this specific role. No generic opener. No "I am writing to apply." Start mid-thought.
2. Evidence: expand on the most relevant project or experience. Specific, concrete, no waffle.
3. Fit: demonstrate genuine knowledge of the company — use the job description and any context provided. One specific observation about their work, product, or team that shows you actually read it.
4. Close: brief, confident. One sentence on why this role specifically, then sign off. No "I look forward to hearing from you."

RULES:
- Never: "I am passionate about", "I am excited to", "I would love to", "keen to learn", "fast-paced environment"
- Match tone to company: early-stage startup = punchy and direct; bank/larger firm = slightly more formal but never stiff
- Do NOT mention any algorithm names or quant jargon from the trading system (no NSGA-III, PPO, DSR, PBO, CPCV) — describe the engineering, not the theory
- Every claim must be defensible in an interview — nothing that invites a question Harry can't answer
- Format: Dear [Hiring Team / Name if provided],\\n\\n[para]\\n\\n[para]\\n\\n[para]\\n\\n[para]\\n\\nYours sincerely,\\nHarry Winter

HARRY'S PROFILE (what is true and what is not):
{_profile_text}"""

COLD_EMAIL_SYSTEM = f"""You write cold emails for Harry Winter seeking a summer 2026 internship at tech and data companies.

STANDARD: The recipient is a founder or data team lead. They get dozens of these. Most get deleted in 3 seconds. This one should not.

FORMAT:
Subject: [21–40 characters, no spam language, no "opportunity", no "internship enquiry" — make it specific to THEIR work]
---
Hi [Name],

[Body: 50–100 words across 3 short paragraphs]

— Harry

SUBJECT LINE FORMULA:
Use one of these patterns (adapt to what's real about this company):
  "[their product/blog/talk] – UCL student question"
  "Your [specific thing] – Harry Winter"
  "[company-specific angle] – data internship"
Never use: "Quick question", "Internship enquiry", "Following up", or anything generic.

BODY STRUCTURE:
Paragraph 1 (1 sentence): Open with THEM, not Harry. Use the formula: "I [action] your [specific content] about [topic], and [genuine insight or reaction]." Must reference something real from the company context provided — a product detail, blog post angle, or technical approach. Never "I admire your mission" or "I came across your company."
Paragraph 2 (1-2 sentences): Who Harry is + one specific credential tied to THEIR problem. The trading system if data/quant/Python. The Svitlo sprint if product/client/teamwork. Not both. Link GitHub when directly relevant.
Paragraph 3 (1 sentence): The ask. Specific and time-bound: "Would you have 15 minutes this week or next — happy to fit your schedule." If the context suggests no open roles, add: "Or if timing isn't right, happy to keep an eye on graduate roles next year."
Close: "— Harry" only. No phone number, no title, no LinkedIn URL.

RULES:
- No attachments mentioned — link to GitHub when relevant, never "please find attached"
- No quant jargon from the trading system (no NSGA-III, PPO, DSR, PBO)
- Never exceed 100 words in the body
- Output format must be: Subject line, then "---", then the email body
- If UCL alumni context is provided, open with the shared connection instead of the company observation

HARRY'S PROFILE:
{_profile_text}"""

LINKEDIN_DM_SYSTEM = f"""You write LinkedIn direct messages for Harry Winter seeking a summer 2026 internship.

STANDARD: A LinkedIn DM is read on a phone. It must work in 10 seconds. If it could be sent to anyone, it fails.

FORMAT: 40–60 words. No greeting beyond first name. No sign-off beyond "— Harry".

STRUCTURE:
Sentence 1: specific observation about their work/company that shows genuine attention
Sentence 2: one-line hook — who Harry is and the single most relevant credential
Sentence 3: a direct, low-friction ask

RULES:
- Start with "Hi [Name]" — nothing else on that line
- No quant jargon
- Never mention "internship programme" or "hiring process" — be human
- Output only the message, nothing else

HARRY'S PROFILE:
{_profile_text}"""

FOLLOWUP_SYSTEM = f"""You write brief follow-up emails for Harry Winter, checking in on a previous cold email.

STANDARD: Under 40 words in the body. Polite, not desperate. Each follow-up must give a NEW reason to reply — not just a bump.

FORMAT:
Subject: Re: [derive from original subject context]
---
Hi [Name],

[Body]

— Harry

FOLLOW-UP STAGE RULES (follow the stage number provided):

Stage 1 (Day 3 — The Bump):
Brief, non-grovelling resurface. 2 sentences max.
- Sentence 1: "Wanted to resurface this in case it got buried."
- Sentence 2: soft re-ask — "Still happy to chat if useful — happy to fit your schedule."

Stage 2 (Day 7 — The Value-Add):
Lead with something Harry built or learned since the original email. This is the most important follow-up.
- Sentence 1: "Since I reached out, I [built / explored / shipped] [specific thing related to their sector]."
- Sentence 2: include a GitHub link if provided, or describe the output concisely.
- Sentence 3: "Happy to share more if it's relevant to how your team handles [their problem]."

Stage 3 (Day 14 — The Breakup):
Final note. Short, gracious, leaves the door open.
- "This is my last note — I completely understand if the timing isn't right."
- "If things change, I'd still love to connect. Thanks for considering."

UNIVERSAL RULES:
- Never: "I was just following up", "I hope I'm not bothering you", "just wondering"
- Body under 40 words
- Output: Subject line, then "---", then the body. Nothing else.

HARRY'S PROFILE:
{_profile_text}"""


# ── API client ─────────────────────────────────────────────────────────────────

_client = None


def _get_client():
    global _client
    if _client is None:
        if not ANTHROPIC_API_KEY:
            raise ValueError("ANTHROPIC_API_KEY not set — add it to your .env file.")
        _client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    return _client


def _call(system_prompt: str, user_message: str, max_tokens: int = 1000) -> tuple[str, dict]:
    """Single API call with prompt caching on system prompt."""
    client = _get_client()
    response = client.messages.create(
        model=GENERATE_MODEL,
        max_tokens=max_tokens,
        system=[
            {
                "type": "text",
                "text": system_prompt,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        messages=[{"role": "user", "content": user_message}],
    )
    usage = {
        "input_tokens": response.usage.input_tokens,
        "output_tokens": response.usage.output_tokens,
        "cache_read": getattr(response.usage, "cache_read_input_tokens", 0),
        "cache_write": getattr(response.usage, "cache_creation_input_tokens", 0),
    }
    text = next((b.text for b in response.content if isinstance(b, TextBlock)), "")
    return text.strip(), usage


# ── Public API ─────────────────────────────────────────────────────────────────

def generate_cover_letter(company: str, role: str, job_description: str, company_context: str = "") -> tuple[str, dict]:
    """Generate a tailored cover letter. Returns (letter_text, usage_stats)."""
    msg = f"""Generate a cover letter for this role:

Company: {company}
Role: {role}
{f'Company context: {company_context}' if company_context else ''}

Job description:
{job_description[:3500]}

Output only the cover letter, starting with 'Dear'."""
    return _call(COVER_LETTER_SYSTEM, msg, max_tokens=1000)


def generate_cold_email(company: str, role_or_area: str, company_context: str, contact_name: str) -> tuple[str, dict]:
    """
    Generate a cold email with subject line.
    Returns (raw_text, usage) where raw_text format is:
        Subject: [subject line]
        ---
        [email body]
    """
    if not contact_name:
        raise ValueError("Cold email requires a named contact. Find the person first.")
    msg = f"""Write a cold email for:

Contact name: {contact_name}
Company: {company}
Target area: {role_or_area}
Company context (use this for the specific opener): {company_context}

Output the subject line and email body in the specified format."""
    return _call(COLD_EMAIL_SYSTEM, msg, max_tokens=300)


def parse_cold_email(raw: str) -> tuple[str, str]:
    """Split raw output into (subject, body)."""
    if "---" in raw:
        parts = raw.split("---", 1)
        subject = parts[0].replace("Subject:", "").strip()
        body = parts[1].strip()
    else:
        subject = ""
        body = raw.strip()
    return subject, body


def generate_followup_email(
    company: str,
    contact_name: str,
    original_subject: str = "",
    days_since: int = 7,
    follow_up_number: int = 1,
    github_link: str = "",
) -> tuple[str, dict]:
    """Generate a follow-up email.
    follow_up_number: 1 = Day 3 bump, 2 = Day 7 value-add, 3 = Day 14 breakup.
    Returns (raw_text, usage) in Subject/---/body format.
    """
    stage_map = {1: "Stage 1 (Day 3 — The Bump)", 2: "Stage 2 (Day 7 — The Value-Add)", 3: "Stage 3 (Day 14 — The Breakup)"}
    stage_label = stage_map.get(follow_up_number, stage_map[1])
    msg = f"""Write a follow-up email — {stage_label}:

Contact name: {contact_name if contact_name and contact_name.lower() != 'there' else 'unknown (use Hi there,)'}
Company: {company}
Days since original email: {days_since}
{f'Original subject line: {original_subject}' if original_subject else ''}
{f'GitHub link to include: {github_link}' if github_link else ''}

Output in the format: Subject line, then "---", then the body."""
    return _call(FOLLOWUP_SYSTEM, msg, max_tokens=200)


def generate_linkedin_dm(company: str, role_or_area: str, company_context: str, contact_name: str = "") -> tuple[str, dict]:
    """Generate a LinkedIn direct message. Returns (message_text, usage_stats)."""
    msg = f"""Write a LinkedIn DM for:

{f'Contact name: {contact_name}' if contact_name else 'Contact: unknown (use Hi there,)'}
Company: {company}
Area of interest: {role_or_area}
Company context: {company_context}

Output only the message."""
    return _call(LINKEDIN_DM_SYSTEM, msg, max_tokens=200)


# ── File output ────────────────────────────────────────────────────────────────

def save_cover_letter_docx(company: str, role: str, letter_text: str) -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Harry Winter\n")
    r.font.size = Pt(11)
    r.font.bold = True
    sub = header.add_run(
        "+44 7453419094  |  harrywinter.uk@gmail.com  |  London, UK\n"
        "linkedin.com/in/harry-p-winter"
    )
    sub.font.size = Pt(9)

    doc.add_paragraph()
    dp = doc.add_paragraph(datetime.now().strftime("%d %B %Y"))
    dp.runs[0].font.size = Pt(10)
    ap = doc.add_paragraph(company)
    ap.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    for para_text in letter_text.split("\n\n"):
        p = doc.add_paragraph(para_text.strip())
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.size = Pt(10.5)

    safe_co = "".join(c for c in company if c.isalnum() or c in " -_").strip().replace(" ", "_")
    safe_role = "".join(c for c in role if c.isalnum() or c in " -_").strip().replace(" ", "_")[:30]
    filename = COVER_LETTER_DIR / f"{datetime.now().strftime('%Y%m%d')}_{safe_co}_{safe_role}.docx"
    doc.save(str(filename))
    return filename


def save_outreach(company: str, subject: str, body: str, mode: str = "email") -> Path:
    safe_co = "".join(c for c in company if c.isalnum() or c in " -_").strip().replace(" ", "_")
    date_str = datetime.now().strftime("%Y%m%d")
    filename = OUTREACH_DIR / f"{date_str}_{safe_co}_{mode}.txt"
    content = f"Subject: {subject}\n\n{body}" if subject else body
    filename.write_text(content, encoding="utf-8")
    return filename


def estimate_cost(usage: dict) -> float:
    """Approximate cost in USD for a Sonnet call."""
    # Sonnet 4.6 pricing: $3/MTok input, $15/MTok output, cache read $0.30/MTok, cache write $3.75/MTok
    normal_input = usage["input_tokens"] - usage.get("cache_read", 0) - usage.get("cache_write", 0)
    cost = (
        normal_input * 3.00 / 1_000_000
        + usage.get("cache_read", 0) * 0.30 / 1_000_000
        + usage.get("cache_write", 0) * 3.75 / 1_000_000
        + usage["output_tokens"] * 15.00 / 1_000_000
    )
    return round(cost, 5)
